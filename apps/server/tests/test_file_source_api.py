from collections.abc import Callable
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from typing import Any, BinaryIO
from uuid import UUID
from zipfile import ZipFile

from database_support import migrated_database
from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter

from liyan_server.app import create_app
from liyan_server.auth import InvalidAccessToken, VerifiedIdentity
from liyan_server.file_parse_worker import process_file_parse
from liyan_server.file_parsing import FileParseLimits
from liyan_server.object_storage import (
    ObjectStorage,
    ObjectStorageState,
    ObjectStorageUnconfigured,
    StoredObject,
)
from liyan_server.settings import Settings


class DeterministicJwtVerifier:
    def verify(self, token: str) -> VerifiedIdentity:
        identities = {
            "allowed-token": VerifiedIdentity(
                subject="supabase-user-1", email="writer@example.com"
            ),
            "second-token": VerifiedIdentity(subject="supabase-user-2", email="second@example.com"),
        }
        try:
            return identities[token]
        except KeyError as error:
            raise InvalidAccessToken from error


class MemoryObjectStorage(ObjectStorage):
    def __init__(self) -> None:
        self.objects: dict[str, bytes] = {}
        self.on_open_read: Callable[[], None] | None = None

    def put(self, key: str, stream: BinaryIO, *, content_type: str) -> None:
        self.objects[key] = stream.read()

    def open(self, key: str) -> BytesIO:
        content = self.objects[key]
        callback = self.on_open_read

        class CallbackBytesIO(BytesIO):
            called = False

            def read(self, size: int | None = -1) -> bytes:
                if callback is not None and not self.called:
                    self.called = True
                    callback()
                return super().read(-1 if size is None else size)

        return CallbackBytesIO(content)

    def delete(self, key: str) -> None:
        self.objects.pop(key, None)

    def list_objects(self, prefix: str = "") -> tuple[StoredObject, ...]:
        return tuple(
            StoredObject(key=key, written_at=datetime.now(UTC))
            for key in sorted(self.objects)
            if key.startswith(prefix)
        )


class RecordingExecutionDispatcher:
    def __init__(self, database_url: str, storage: ObjectStorage) -> None:
        self.database_url = database_url
        self.storage = storage
        self.execution_ids: list[UUID] = []
        self.limits = FileParseLimits(
            max_pages=20,
            max_normalized_characters=10_000,
            timeout_seconds=10,
            max_docx_entries=100,
            max_docx_uncompressed_bytes=1_000_000,
        )

    def dispatch(self, execution_id: UUID) -> None:
        self.execution_ids.append(execution_id)

    def is_reachable(self) -> bool:
        return True

    def run_next(self) -> None:
        process_file_parse(
            self.database_url,
            self.execution_ids.pop(0),
            self.storage,
            limits=self.limits,
            short_source_characters=20,
        )


def authenticated_client(
    tmp_path: Path, *, max_file_bytes: int = 1_000_000
) -> tuple[TestClient, dict[str, str], RecordingExecutionDispatcher, MemoryObjectStorage]:
    database_url = migrated_database(tmp_path)
    storage = MemoryObjectStorage()
    dispatcher = RecordingExecutionDispatcher(database_url, storage)
    settings = Settings(
        database_url=database_url,
        allowed_emails="writer@example.com,second@example.com",
        file_max_bytes=max_file_bytes,
    )
    client = TestClient(
        create_app(
            settings,
            jwt_verifier=DeterministicJwtVerifier(),
            execution_dispatcher=dispatcher,
            object_storage=storage,
        )
    )
    return client, {"Authorization": "Bearer allowed-token"}, dispatcher, storage


def test_uploading_text_streams_an_owner_scoped_object_and_creates_parse_execution(
    tmp_path: Path,
) -> None:
    client, headers, dispatcher, storage = authenticated_client(tmp_path)

    response = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "source-1"},
        files={"file": ("notes.md", b"# Heading\n\nUseful body.", "text/markdown")},
    )

    assert response.status_code == 201
    source = response.json()
    assert source["filename"] == "notes.md"
    assert source["content_type"] == "text/markdown"
    assert source["size_bytes"] == 23
    assert source["status"] == "processing"
    assert source["title"] is None
    assert source["body"] is None
    assert source["active_execution"]["operation"] == "parse_file"
    assert source["active_execution"]["status"] == "queued"
    assert dispatcher.execution_ids == [UUID(source["active_execution"]["id"])]
    assert len(storage.objects) == 1
    object_key = next(iter(storage.objects))
    assert object_key.startswith("users/")
    assert "/source-preparations/" in object_key
    assert storage.objects[object_key] == b"# Heading\n\nUseful body."


def test_text_parse_result_is_editable_and_owner_isolated(tmp_path: Path) -> None:
    client, headers, dispatcher, _ = authenticated_client(tmp_path)
    created = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "source-1"},
        files={"file": ("notes.txt", b"First line\r\n\r\nUseful body.", "text/plain")},
    ).json()

    dispatcher.run_next()
    response = client.get(f"/task-creation/file-sources/{created['id']}", headers=headers)

    assert response.status_code == 200
    source = response.json()
    assert source["status"] == "ready"
    assert source["title"] == "notes"
    assert source["body"] == "First line\n\nUseful body."
    assert source["provenance"] == "notes.txt"
    assert source["active_execution"]["status"] == "succeeded"
    edited = client.patch(
        f"/task-creation/file-sources/{created['id']}/content",
        headers=headers,
        json={
            "title": "Edited notes",
            "body": "Edited body.",
            "provenance": "Notebook",
        },
    )
    assert edited.status_code == 200
    assert edited.json()["input_version"] == 2
    assert edited.json()["title"] == "Edited notes"
    assert (
        client.get(
            f"/task-creation/file-sources/{created['id']}",
            headers={"Authorization": "Bearer second-token"},
        ).status_code
        == 404
    )


def test_declared_and_actual_file_type_must_agree_before_storage(tmp_path: Path) -> None:
    client, headers, dispatcher, storage = authenticated_client(tmp_path)

    response = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "source-1"},
        files={"file": ("report.pdf", b"plain text", "application/pdf")},
    )

    assert response.status_code == 422
    assert response.json()["detail"] == "The uploaded file does not match its declared type."
    assert dispatcher.execution_ids == []
    assert storage.objects == {}


def test_file_size_limit_is_enforced_while_streaming(tmp_path: Path) -> None:
    client, headers, dispatcher, storage = authenticated_client(tmp_path, max_file_bytes=8)

    response = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "source-1"},
        files={"file": ("notes.txt", b"too many bytes", "text/plain")},
    )

    assert response.status_code == 413
    assert dispatcher.execution_ids == []
    assert storage.objects == {}


def test_docx_is_parsed_without_an_llm(tmp_path: Path) -> None:
    client, headers, dispatcher, _ = authenticated_client(tmp_path)
    document = Document()
    document.add_heading("Document heading", level=1)
    document.add_paragraph("A deterministic DOCX paragraph with useful content.")
    content = BytesIO()
    document.save(content)

    created = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "source-1"},
        files={
            "file": (
                "report.docx",
                content.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    ).json()
    dispatcher.run_next()

    source = client.get(f"/task-creation/file-sources/{created['id']}", headers=headers).json()
    assert source["status"] == "ready"
    assert source["title"] == "report"
    assert source["body"] == (
        "Document heading\nA deterministic DOCX paragraph with useful content."
    )


def test_scanned_pdf_fails_only_its_source_with_safe_reason(tmp_path: Path) -> None:
    client, headers, dispatcher, storage = authenticated_client(tmp_path)
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    content = BytesIO()
    writer.write(content)
    created = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "source-1"},
        files={"file": ("scan.pdf", content.getvalue(), "application/pdf")},
    ).json()

    dispatcher.run_next()
    source = client.get(f"/task-creation/file-sources/{created['id']}", headers=headers).json()

    assert source["status"] == "failure"
    assert source["failure"]["code"] == "scanned_document"
    assert "OCR" in source["failure"]["message"]
    assert source["active_execution"]["status"] == "failed"
    assert len(storage.objects) == 1


def test_empty_and_oversized_normalized_text_fail_at_the_parse_boundary(
    tmp_path: Path,
) -> None:
    client, headers, dispatcher, _ = authenticated_client(tmp_path)
    empty = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "empty"},
        files={"file": ("empty.txt", b"   \n", "text/plain")},
    ).json()
    large = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "large"},
        files={"file": ("large.md", b"x" * 10_001, "text/markdown")},
    ).json()

    dispatcher.run_next()
    dispatcher.run_next()

    empty_result = client.get(f"/task-creation/file-sources/{empty['id']}", headers=headers).json()
    large_result = client.get(f"/task-creation/file-sources/{large['id']}", headers=headers).json()
    assert empty_result["failure"]["code"] == "empty_document"
    assert large_result["failure"]["code"] == "normalized_text_too_large"


def test_cancelling_a_running_parse_prevents_late_result_acceptance(tmp_path: Path) -> None:
    client, headers, dispatcher, storage = authenticated_client(tmp_path)
    created = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "source-1"},
        files={"file": ("notes.txt", b"Useful body after cancellation.", "text/plain")},
    ).json()
    execution_id = created["active_execution"]["id"]

    def cancel() -> None:
        response = client.post(f"/executions/{execution_id}/cancel", headers=headers)
        assert response.status_code == 202
        assert response.json()["status"] == "cancel_requested"

    storage.on_open_read = cancel
    dispatcher.run_next()

    source = client.get(f"/task-creation/file-sources/{created['id']}", headers=headers).json()
    assert source["status"] == "failure"
    assert source["failure"]["code"] == "cancelled"
    assert source["title"] is None
    assert source["body"] is None
    assert source["active_execution"]["status"] == "cancelled"
    assert source["active_execution"]["result_id"] is not None


def test_encrypted_and_damaged_docx_files_get_source_specific_failures(tmp_path: Path) -> None:
    client, headers, dispatcher, _ = authenticated_client(tmp_path)
    encrypted = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "encrypted"},
        files={
            "file": (
                "encrypted.docx",
                bytes.fromhex("D0CF11E0A1B11AE1") + b"encrypted payload",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    ).json()
    damaged_content = BytesIO()
    with ZipFile(damaged_content, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
    damaged = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "damaged"},
        files={
            "file": (
                "damaged.docx",
                damaged_content.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    ).json()
    truncated = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "truncated"},
        files={
            "file": (
                "truncated.docx",
                b"PK\x03\x04truncated archive",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    ).json()

    dispatcher.run_next()
    dispatcher.run_next()
    dispatcher.run_next()

    encrypted_result = client.get(
        f"/task-creation/file-sources/{encrypted['id']}", headers=headers
    ).json()
    damaged_result = client.get(
        f"/task-creation/file-sources/{damaged['id']}", headers=headers
    ).json()
    truncated_result = client.get(
        f"/task-creation/file-sources/{truncated['id']}", headers=headers
    ).json()
    assert encrypted_result["failure"]["code"] == "encrypted_document"
    assert damaged_result["failure"]["code"] == "damaged_document"
    assert truncated_result["failure"]["code"] == "damaged_document"


def test_zero_parse_timeout_fails_the_source_before_accepting_content(tmp_path: Path) -> None:
    client, headers, dispatcher, _ = authenticated_client(tmp_path)
    dispatcher.limits = FileParseLimits(
        max_pages=20,
        max_normalized_characters=10_000,
        timeout_seconds=0,
        max_docx_entries=100,
        max_docx_uncompressed_bytes=1_000_000,
    )
    created = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "source-1"},
        files={"file": ("notes.txt", b"Useful content.", "text/plain")},
    ).json()

    dispatcher.run_next()

    source = client.get(f"/task-creation/file-sources/{created['id']}", headers=headers).json()
    assert source["failure"]["code"] == "parse_timeout"
    assert source["title"] is None
    assert source["body"] is None


def test_prepared_file_can_be_replaced_without_losing_its_session_identity(
    tmp_path: Path,
) -> None:
    client, headers, dispatcher, _ = authenticated_client(tmp_path)
    created = client.post(
        "/task-creation/file-sources",
        headers=headers,
        data={"client_session_id": "session-1", "client_source_id": "source-1"},
        files={"file": ("old.txt", b"Old file body.", "text/plain")},
    ).json()
    dispatcher.run_next()

    replaced = client.put(
        f"/task-creation/file-sources/{created['id']}",
        headers=headers,
        files={"file": ("new.md", b"# New\n\nReplacement body.", "text/markdown")},
    )

    assert replaced.status_code == 202
    assert replaced.json()["id"] == created["id"]
    assert replaced.json()["client_source_id"] == "source-1"
    assert replaced.json()["filename"] == "new.md"
    assert replaced.json()["input_version"] == 2
    assert replaced.json()["status"] == "processing"
    assert replaced.json()["active_execution"]["attempt"] == 1
    dispatcher.run_next()
    ready = client.get(
        f"/task-creation/file-sources/{created['id']}", headers=headers
    ).json()
    assert ready["status"] == "ready"
    assert ready["title"] == "new"
    assert ready["body"] == "# New\n\nReplacement body."


class UnconfiguredStorage(MemoryObjectStorage):
    """An operator never filled in LIYAN_R2_*."""

    def missing_settings(self) -> tuple[str, ...]:
        return ("LIYAN_R2_BUCKET",)

    def put(self, key: str, stream: BinaryIO, *, content_type: str) -> None:
        raise ObjectStorageUnconfigured("Object storage is not configured: LIYAN_R2_BUCKET.")

    def state(self) -> ObjectStorageState:
        return "unconfigured"


class UnreachableStorage(MemoryObjectStorage):
    """A configured bucket having a bad minute."""

    def put(self, key: str, stream: BinaryIO, *, content_type: str) -> None:
        raise OSError("Connection reset by peer.")


def _upload_to(tmp_path: Path, storage: MemoryObjectStorage) -> Any:
    database_url = migrated_database(tmp_path)
    dispatcher = RecordingExecutionDispatcher(database_url, storage)
    client = TestClient(
        create_app(
            Settings(database_url=database_url, allowed_emails="writer@example.com"),
            jwt_verifier=DeterministicJwtVerifier(),
            execution_dispatcher=dispatcher,
            object_storage=storage,
        )
    )
    return client.post(
        "/task-creation/file-sources",
        headers={"Authorization": "Bearer allowed-token"},
        data={"client_session_id": "session-1", "client_source_id": "source-1"},
        files={"file": ("notes.md", b"# Heading\n\nUseful body.", "text/markdown")},
    )


def test_an_upload_without_configured_storage_does_not_advise_a_pointless_retry(
    tmp_path: Path,
) -> None:
    response = _upload_to(tmp_path, UnconfiguredStorage())

    assert response.status_code == 503
    detail = response.json()["detail"]
    assert "not configured" in detail
    assert "Retrying will not help" in detail
    # The two working intake routes never touch object storage, so the answer
    # can honestly point at them.
    assert "Paste the text or submit a URL instead" in detail


def test_a_transient_storage_fault_still_reads_as_worth_retrying(tmp_path: Path) -> None:
    response = _upload_to(tmp_path, UnreachableStorage())

    assert response.status_code == 503
    assert response.json()["detail"] == "The file could not be stored. Try again later."


def test_the_two_storage_failures_never_give_the_same_answer(tmp_path: Path) -> None:
    unconfigured = _upload_to(tmp_path, UnconfiguredStorage())
    unreachable = _upload_to(tmp_path, UnreachableStorage())

    assert unconfigured.json()["detail"] != unreachable.json()["detail"]
